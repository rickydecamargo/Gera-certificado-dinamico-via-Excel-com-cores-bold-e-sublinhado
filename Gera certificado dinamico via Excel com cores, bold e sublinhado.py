#Instalar o python-docx
#Instalar o openpyxl
#Precisa instalar o pywin32

#Cria diversos certificados de forma dinâmica baseada em um arquivo excel e um arquivo word já existentes.
#Com base em informações do Excel, alteramos diversas informações base para novas vindas do xlsx e adiciona o nome do curso com cor, bold e underline.
#Após criar os arquivos envia por email. Os email são pegos da planilha Excel DadosAlunosEmail.

from docx import Document
from docx.shared import Pt #Aumenta o tamanho da fonte
from docx.shared import RGBColor #Mudar a cor da fonte
from openpyxl import load_workbook #Abre o arquivo de Excel

import win32com.client as win32
outlook = win32.Dispatch("outlook.application")

#Variavel indicando o arquivo Alunos.xlsx
nome_arquivo_alunos = r"C:\Users\Windows\Desktop\Python Projetos\Word\DadosAlunosEmail.xlsx"

#Abrindo o arquivo
planilhaDadosAluno = load_workbook(nome_arquivo_alunos)

#Selecionando a aba Nomes
sheet_selecionada = planilhaDadosAluno["Nomes"]

#Para realizar a leitura de cada linha no arquivo excel.
for linha in range(2, len(sheet_selecionada["A"]) + 1):

    #Abrindo o arquivo do Word
    arquivoWord = Document("C:\\Users\\Windows\\Desktop\\Python Projetos\\Word\Certificado2.docx")

    #configurando os estilos
    estilo = arquivoWord.styles["Normal"]

    nomeAluno = sheet_selecionada['A%s' % linha].value #Pega o nome de forma dinâmica
    dia = sheet_selecionada['B%s' % linha].value  # Pega o nome de forma dinâmica
    mes = sheet_selecionada['C%s' % linha].value  # Pega o nome de forma dinâmica
    ano = sheet_selecionada['D%s' % linha].value  # Pega o nome de forma dinâmica
    nomeCurso = sheet_selecionada['E%s' % linha].value  # Pega o nome de forma dinâmica
    nomeInstrutor = sheet_selecionada['F%s' % linha].value  # Pega o nome de forma
    emailAluno = sheet_selecionada['G%s' % linha].value  # Pega o nome de forma dinâmica


    frase_parte1 = "Concluiu com sucesso o curso de "
    frase_parte2 = ", com a carga horária de 20 horas, promovido pela escola de Cursos Online em "
    frase_montada = f"{frase_parte2} {dia} de {mes} de {ano} "

    #for = para
    for paragrafo in arquivoWord.paragraphs:

        #Se existir @nome no paragrafo substituir pelo nome indicado.
        if "@nome" in paragrafo.text:
            paragrafo.text = nomeAluno #Utiliza esta variável para alterar os nome que vem do Excel.
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)" #Selecionar a fonte
            fonte.size = Pt(24) #Para mudar o tamanho da fonte

        if "Dezembro" in paragrafo.text:
            paragrafo.text = frase_parte1 #Utiliza esta variável para alterar o texto que armazenamos na variável.
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)" #Selecionar a fonte
            fonte.size = Pt(24) #Para mudar o tamanho da fonte
            adicionaNovaPalavra = paragrafo.add_run(nomeCurso) #Adiciona o nome do curso no final do paragrafo
            adicionaNovaPalavra.font.color.rgb = RGBColor(255,0,0) #Muda a cor para Vermelho
            adicionaNovaPalavra.underline = True #Sublinha a palavra
            adicionaNovaPalavra.bold = True #Negrito
            adicionaNovaPalavra = paragrafo.add_run(frase_montada) #Adiciona o texto final do paragrafo
            adicionaNovaPalavra.font.color.rgb = RGBColor(0,0,0) #Muda a cor para Vermelho


        if "Instrutor" in paragrafo.text:
            paragrafo.text = nomeInstrutor + " - Instrutor(a)" #Utiliza esta variável para alterar os nome que vem do Excel.
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)" #Selecionar a fonte
            fonte.size = Pt(24) #Para mudar o tamanho da fonte

    caminhoCertificado = "C:\\Users\\Windows\\Desktop\\Python Projetos\\Word\\Certificados\\" + nomeAluno + ".docx"

    #salva o certificado com o nome do aluno
    arquivoWord.save(caminhoCertificado)


    #Configurando o disparo dos email com certificados.
    emailOutlook = outlook.CreateItem(0)

    #Pegando o primeiro nome
    primeiroNome = nomeAluno.split(None, 1)[0]

    # Para quem irá enviar o email
    emailOutlook.To = emailAluno
    #Assunto do Email
    emailOutlook.Subject = "Certificado " + nomeAluno
    #Corpo do email dinâmico
    emailOutlook.HTMLBody = f"""
        <p>Boa noite {primeiroNome}.</p>
        <p>Segue seu <b>certificado</b>.</p>
        <p>Atenciosamente.</p>
        <p><img src="C:\\Users\\Windows\\Desktop\\Python Projetos\\Word\Assinatura.jpg"></p>
    """

    #Adicionando o certificado em anexo
    emailOutlook.Attachments.Add(caminhoCertificado)
    emailOutlook.save() #save = cria e salva com draft - send envia o email direto
print("Certificados gerados com sucesso!")